import pdfplumber
import docx
import os
import logging

logger = logging.getLogger(__name__)

class ContractParser:
    def parse_file(self, file_path):
        """
        Parses PDF or DOCX file and returns extracted text.
        """
        ext = os.path.splitext(file_path)[1].lower()
        if ext == '.pdf':
            return self._parse_pdf(file_path)
        elif ext == '.docx':
            return self._parse_docx(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")

    def _parse_pdf(self, file_path):
        text = ""
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            raise
        return text

    def _parse_docx(self, file_path):
        text = ""
        try:
            doc = docx.Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            raise
        return text

    def parse_uploaded_file(self, uploaded_file):
        """
        Handle Streamlit UploadedFile object.
        """
        if uploaded_file.name.endswith('.pdf'):
            with pdfplumber.open(uploaded_file) as pdf:
                text = "\n".join([page.extract_text() or "" for page in pdf.pages])
            return text
        elif uploaded_file.name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            text = "\n".join([para.text for para in doc.paragraphs])
            return text
        else:
            raise ValueError("Unsupported file type")
